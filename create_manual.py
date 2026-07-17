from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).with_name('MANUAL_CALCULOS.docx')


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        shade(table.rows[0].cells[i], '1F4E78')
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.7)
sec.bottom_margin = Cm(1.7)
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)
styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10)
styles['Title'].font.color.rgb = RGBColor(31, 78, 121)
styles['Heading 1'].font.color.rgb = RGBColor(31, 78, 121)
styles['Heading 2'].font.color.rgb = RGBColor(37, 99, 235)

title = doc.add_paragraph(style='Title')
title.add_run('Validador de planificaciones')
sub = doc.add_paragraph('Manual funcional y metodología de cálculo')
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(37, 99, 235)

p = doc.add_paragraph()
p.add_run('Versión: ').bold = True
p.add_run('2.2 — balance de franjas y filtro semanal global')

doc.add_heading('1. Selección del origen de horarios', level=1)
doc.add_paragraph('La aplicación puede analizar tres colecciones distintas de horarios contenidas en dayTimes. El usuario selecciona una única fuente y todo el motor se ejecuta exclusivamente sobre ella. Las fuentes no se suman, concatenan ni comparan de forma implícita.')
add_table(doc, ['Clave JSON', 'Nombre mostrado', 'Uso'], [
    ['planned', 'Plan publicado', 'Horario definitivo o publicado.'],
    ['plannedDraft', 'Borrador del planificador', 'Propuesta generada por el planificador.'],
    ['plannedDraftManuallyEdited', 'Borrador editado manualmente', 'Versión modificada manualmente.'],
])

doc.add_heading('2. Flujo de selección y filtro semanal', level=1)
for text in [
    'El usuario carga un JSON o TXT.',
    'La aplicación inspecciona planned, plannedDraft y plannedDraftManuallyEdited.',
    'Solo aparecen en el selector las fuentes que contienen al menos un segmento WORK.',
    'El selector Semana a analizar permite mantener todo el periodo o elegir una semana ISO.',
    'El filtro semanal afecta a turnos, horas contractuales, cobertura diaria, balance mañana/tarde, ausencias e incidencias.',
    'Los resúmenes cuya unidad original es empleado-mes mantienen su periodo mensual y se identifican como tales.',
]:
    doc.add_paragraph(text, style='List Number')

doc.add_heading('3. Formación del turno', level=1)
doc.add_paragraph('Dentro de la fuente seleccionada solo se consideran segmentos con hourType = WORK. Para cada empleado y fecha se forma un único turno diario.')
add_table(doc, ['Medida', 'Cálculo'], [
    ['Inicio', 'Mínimo startDateTime de los segmentos WORK seleccionados.'],
    ['Fin', 'Máximo endDateTime de los segmentos WORK seleccionados.'],
    ['Horas trabajadas', 'Suma de endDateTime - startDateTime de todos los segmentos WORK.'],
    ['Descanso interno', 'Suma de huecos positivos de hasta una hora entre segmentos consecutivos.'],
])

doc.add_heading('4. Restricciones', level=1)
add_table(doc, ['Regla', 'Incidencia'], [
    ['Duración máxima', 'Horas trabajadas > 7,5.'],
    ['Duración mínima', 'Horas trabajadas < 4,0.'],
    ['Descanso entre jornadas', 'Inicio del turno actual - fin del anterior < 11 horas.'],
    ['Días consecutivos', 'Racha de fechas trabajadas > 5 días.'],
])
doc.add_paragraph('Las cuatro reglas son idénticas para los tres orígenes. Al seleccionar una semana, el detalle de incidencias se limita a las que intersectan ese intervalo.')

doc.add_heading('5. Control semanal de horas', level=1)
doc.add_paragraph('Las horas planificadas semanales son la suma de horas trabajadas de la fuente seleccionada por empleado y semana ISO.')
add_table(doc, ['Medida', 'Fórmula'], [
    ['Diferencia', 'horas planificadas - applicableWorkingHours'],
    ['Horas faltantes', 'max(applicableWorkingHours - horas planificadas, 0)'],
    ['Horas en exceso', 'max(horas planificadas - applicableWorkingHours, 0)'],
    ['COINCIDE', '|diferencia| <= 0,01 horas en semana completa.'],
    ['NO EVALUABLE', 'La semana no contiene los siete días del fichero.'],
    ['SIN HORAS CONTRATO', 'applicableWorkingHours no es numérico.'],
])
doc.add_paragraph('La tabla de detalle incluye horas teóricas asociadas a ausencias, horas planificadas más ausencia estimada y una etiqueta que indica si la ausencia podría explicar todo o parte del déficit. Esta información es diagnóstica y no imputa horas trabajadas.')

doc.add_heading('6. Balance de turnos de mañana y tarde', level=1)
doc.add_paragraph('La franja se determina exclusivamente por la hora de inicio del turno diario.')
add_table(doc, ['Concepto', 'Definición'], [
    ['Turno de mañana', 'Inicio anterior a las 13:00.'],
    ['Turno de tarde', 'Inicio a las 13:00 o posterior.'],
    ['Solo mañanas', 'El empleado no tiene ningún turno de tarde en el periodo filtrado.'],
    ['Solo tardes', 'El empleado no tiene ningún turno de mañana en el periodo filtrado.'],
    ['Índice de equilibrio', '2 × min(turnos de mañana, turnos de tarde) / turnos totales × 100.'],
    ['Mañanas medias por semana', 'Turnos de mañana / semanas incluidas en el periodo seleccionado.'],
    ['Tardes medias por semana', 'Turnos de tarde / semanas incluidas en el periodo seleccionado.'],
])
doc.add_paragraph('La gráfica principal compara dos columnas por empleado: mañanas medias por semana y tardes medias por semana. Si se selecciona una sola semana, el denominador es uno.')

doc.add_heading('7. Ausencias', level=1)
doc.add_paragraph('Las ausencias se leen siempre desde dayTimes.absences, independientemente del origen horario. Solo se incluyen estados VALIDATED o APPROVED. Las horas potenciales asociadas a ausencia no se suman a la planificación; sirven únicamente como explicación del déficit.')
doc.add_paragraph('El calendario diario cuenta empleados únicos ausentes por fecha y conserva explícitamente los días con cero ausencias. El filtro semanal limita el calendario al intervalo seleccionado.')

doc.add_heading('8. Indicadores del dashboard', level=1)
add_table(doc, ['Indicador', 'Unidad y cálculo'], [
    ['Empleados', 'Personas únicas dentro del ámbito de la visualización.'],
    ['Turnos', 'Registros empleado-día con al menos un segmento WORK.'],
    ['Horas planificadas', 'Suma neta de horas de los turnos seleccionados.'],
    ['Sin incidencias', 'Empleado que cumple las cuatro reglas en el periodo mensual evaluado.'],
    ['Incidencias', 'Número de excepciones que intersectan el periodo filtrado.'],
    ['Cumplimiento semanal', 'Registros COINCIDE / registros semanales evaluables.'],
    ['Rotan mañana y tarde', 'Empleados con al menos un turno en ambas franjas.'],
    ['Mañanas/tardes medias', 'Número de turnos de cada franja dividido por las semanas del ámbito.'],
    ['Empleados ausentes por día', 'Personas únicas con ausencia validada o aprobada en cada fecha.'],
])

doc.add_heading('9. Casos y advertencias', level=1)
for text in [
    'Las fuentes nunca se mezclan. El dashboard representa una sola selección cada vez.',
    'El filtro semanal no convierte métricas mensuales en métricas semanales.',
    'Una semana parcial puede aparecer como no evaluable en el control contractual.',
    'Los segmentos solapados se suman y pueden duplicar tiempo; deben corregirse en origen.',
    'El promedio de mañanas y tardes usa como denominador todas las semanas incluidas en el ámbito seleccionado.',
    'Las ausencias son una explicación potencial del déficit, no una prueba causal ni horas trabajadas.',
]:
    doc.add_paragraph(text, style='List Bullet')

doc.add_heading('10. Uso', level=1)
for text in [
    'Ejecutar lanzar_app.bat.',
    'Subir el fichero JSON o TXT.',
    'Seleccionar el origen de horarios en el panel lateral.',
    'Seleccionar todas las semanas o una semana ISO concreta.',
    'Comprobar en la barra lateral el periodo activo.',
    'Analizar las pestañas o descargar el Excel de detalle.',
]:
    doc.add_paragraph(text, style='List Number')

doc.save(OUT)
print(OUT)
